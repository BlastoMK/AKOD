import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import openpyxl
import mysql.connector
from datetime import datetime
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

current_excel_file = r"C:\Users\blast\OneDrive\Desktop\Demo\AKOD\RMU Excel.xlsx"

def open_excel_data_entry():
    def load_data():
        try:
            if current_excel_file and os.path.exists(current_excel_file):
                workbook = openpyxl.load_workbook(current_excel_file)
                sheet = workbook.active
                list_values = list(sheet.values)
                print(list_values)
            else:
                messagebox.showerror("Error", "No Excel file found at the specified path")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while loading data: {e}")

    def calculate_days(date1_str, date2_str):
        date_format = "%d/%m/%Y"
        date1 = datetime.strptime(date1_str, date_format)
        date2 = datetime.strptime(date2_str, date_format)
        delta = date2 - date1
        return delta.days

    def update_calculations(event):
        try:
            date_received = entry_widgets[6].get()
            date_marked = entry_widgets[10].get()
            response_date = entry_widgets[12].get()

            if date_received and date_marked:
                days_taken_to_mark = calculate_days(date_received, date_marked)
                entry_widgets[11].config(state='normal')
                entry_widgets[11].delete(0, 'end')
                entry_widgets[11].insert(0, str(days_taken_to_mark))
                entry_widgets[11].config(state='readonly')
            if date_received and response_date:
                days_taken_to_respond = calculate_days(date_received, response_date)
                entry_widgets[13].config(state='normal')
                entry_widgets[13].delete(0, 'end')
                entry_widgets[13].insert(0, str(days_taken_to_respond))
                entry_widgets[13].config(state='readonly')
        except Exception as e:
            print(f"Error in calculation: {e}")

    def insert_row(entry_widgets):
        try:
            global current_excel_file
            row_values = [widget.get() for widget in entry_widgets]

            date_received = row_values[6]
            date_marked = row_values[10]
            response_date = row_values[12]

            days_taken_to_mark = calculate_days(date_received, date_marked)
            days_taken_to_respond = calculate_days(date_received, response_date)

            row_values[11] = str(days_taken_to_mark)
            row_values[13] = str(days_taken_to_respond)

            if current_excel_file and os.path.exists(current_excel_file):
                workbook = openpyxl.load_workbook(current_excel_file)
                sheet = workbook.active
                sheet.append(row_values)
                workbook.save(current_excel_file)

                connection = mysql.connector.connect(
                    host="localhost",
                    user="root",
                    password="Sage_mode2",
                    database="rmu_excel"
                )
                cursor = connection.cursor()
                sql = """INSERT INTO excel_document 
                         (number, receiver, ref, origin, date_on_letter, file_no, date_received, subject, 
                          processing_officer, action_officer, date_marked, days_taken_to_mark, response_date, 
                          days_taken_to_respond, comments, date_of_dispatch) 
                         VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                cursor.execute(sql, row_values)
                connection.commit()
                cursor.close()
                connection.close()

                for widget in entry_widgets:
                    widget.delete(0, 'end')

                messagebox.showinfo("Success", "Data inserted successfully")
            else:
                messagebox.showerror("Error", "No Excel file found at the specified path")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    def generate_report():
        try:
            file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
            if not file_path:
                messagebox.showerror("Error", "No Excel file selected")
                return
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheet = workbook.active
            
            total_mail_received = 0
            mail_for_information = 0
            number_eligible_for_response = 0
            mail_responded_within_10_days = 0
            mail_responded_after_10_days = 0
            mail_not_responded_to = 0
            on_going_mail = 0
            compliance_level = 0

            for row in sheet.iter_rows(min_row=2, values_only=True):
                try:
                    if len(row) > 6 and row[6]:
                        total_mail_received += 1
                        if len(row) > 7 and row[7] and "information" in row[7].lower():
                            mail_for_information += 1
                        else:
                            number_eligible_for_response += 1
                            if len(row) > 11 and isinstance(row[11], (int, float)):
                                days_taken_to_mark = row[11]
                                if days_taken_to_mark <= 10:
                                    mail_responded_within_10_days += 1
                                elif days_taken_to_mark > 10:
                                    mail_responded_after_10_days += 1
                            if len(row) > 13 and not row[13]:
                                on_going_mail += 1
                except Exception as e:
                    print(f"Error parsing row: {e}")

            mail_not_responded_to = number_eligible_for_response - (mail_responded_within_10_days + mail_responded_after_10_days)
            compliance_level = (mail_responded_within_10_days / number_eligible_for_response) * 100 if number_eligible_for_response > 0 else 0

            current_month = datetime.now().strftime("%B")
            report_filename = f"{current_month}_Monthly_Report.docx"
            report_path = os.path.join(os.path.expanduser("~"), "Desktop", report_filename)
            doc = docx.Document()
            
            doc.add_heading('MINISTRY OF TRANSPORT AND PUBLIC WORKS', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('Response to correspondence to 10 days', level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('1st Quarter 2023/24 CORPORATE SERVICES', level=3).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            table = doc.add_table(rows=9, cols=8)
            table.style = 'Table Grid'

            headers = [' ', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
            data = [
                ['Total Mail Received', total_mail_received, '', '', '', '', '', '', ''],
                ['Mail for information', mail_for_information, '', '', '', '', '', '', ''],
                ['Number eligible for response', number_eligible_for_response, '', '', '', '', '', '', '', ''],
                ['Mail Responded to within 10 W/days', mail_responded_within_10_days, '', '', '', '', '', '', '', ''],
                ['Mail Responded after 10 W/days', mail_responded_after_10_days, '', '', '', '', '', '', '', ''],
                ['Mail not Responded to', mail_not_responded_to, '', '', '', '', '', '', '', ''],
                ['On-going (Mail still to be responded to but within time)', on_going_mail, '', '', '', '', '', '', '', ''],
                ['Compliance Level', f"{compliance_level:.1f}%", '', '', '', '', '', '', ''],
            ]

            for i, header in enumerate(headers):
                table.cell(0, i).text = header

            for row_idx, row_data in enumerate(data, start=1):
                for col_idx, value in enumerate(row_data):
                    table.cell(row_idx, col_idx).text = str(value)

            doc.save(report_path)

            messagebox.showinfo("Success", f"Report generated and saved at {report_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while generating the report: {e}")

    def create_new_file():
        try:
            global current_excel_file
            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if file_path:
                new_workbook = openpyxl.Workbook()
                new_sheet = new_workbook.active
                
                if current_excel_file and os.path.exists(current_excel_file):
                    workbook = openpyxl.load_workbook(current_excel_file)
                    sheet = workbook.active
                    headers = [cell.value for cell in sheet[1]]
                    new_sheet.append(headers)
                
                new_workbook.save(file_path)
                current_excel_file = file_path
                
                messagebox.showinfo("Success", "New file created successfully")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while creating a new file: {e}")

    def select_excel_file():
        global current_excel_file
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            current_excel_file = file_path
            messagebox.showinfo("Success", f"Selected file: {current_excel_file}")

    root = tk.Tk()
    root.title("RMU Excel Data Entry")

    style = ttk.Style(root)
    style.theme_use("clam")

    frame = ttk.Frame(root)
    frame.pack(pady=20, padx=20)

    widgets_frame = ttk.LabelFrame(frame, text="Insert Row")
    widgets_frame.grid(row=0, column=0, padx=20, pady=10)

    fields = ["Number", "Receiver", "REF", "Origin", "Date on Letter (DD/MM/YYYY)", "File No",
              "Date Received (DD/MM/YYYY)", "Subject", "Processing Officer", "Action Officer", 
              "Date Marked (DD/MM/YYYY)", "Days Taken to Mark", "Response Date (DD/MM/YYYY)", 
              "Days Taken to Respond", "Comments", "Date of Dispatch (DD/MM/YYYY)"]

    origin_options = ["ATTORNEY", "DBWM", "DCEC", "DFM", "DID", "DRTS", "EB", "GFM", "IE", "ROADS"]

    entry_widgets = []

    for i, field in enumerate(fields):
        label = ttk.Label(widgets_frame, text=field)
        label.grid(row=i, column=0, padx=5, pady=5, sticky="e")
        if field == "Origin":
            combobox = ttk.Combobox(widgets_frame, values=origin_options)
            combobox.grid(row=i, column=1, padx=5, pady=5)
            entry_widgets.append(combobox)
        elif field in ["Days Taken to Mark", "Days Taken to Respond"]:
            entry = ttk.Entry(widgets_frame, state='readonly')
            entry.grid(row=i, column=1, padx=5, pady=5)
            entry_widgets.append(entry)
        else:
            entry = ttk.Entry(widgets_frame)
            entry.grid(row=i, column=1, padx=5, pady=5)
            entry_widgets.append(entry)

    date_fields = [6, 10, 12]
    for idx in date_fields:
        entry_widgets[idx].bind("<FocusOut>", update_calculations)

    insert_button = ttk.Button(widgets_frame, text="Insert", command=lambda: insert_row(entry_widgets))
    insert_button.grid(row=len(fields), column=1, padx=5, pady=10, sticky="ew")

    buttons_frame = ttk.LabelFrame(frame, text="Actions")
    buttons_frame.grid(row=1, column=0, padx=20, pady=10)

    select_file_button = ttk.Button(buttons_frame, text="Select Excel File", command=select_excel_file)
    select_file_button.grid(row=0, column=0, padx=5, pady=10, sticky="ew")

    new_file_button = ttk.Button(buttons_frame, text="New File", command=create_new_file)
    new_file_button.grid(row=1, column=0, padx=5, pady=10, sticky="ew")

    generate_report_button = ttk.Button(buttons_frame, text="Generate Report", command=generate_report)
    generate_report_button.grid(row=2, column=0, padx=5, pady=10, sticky="ew")

    root.mainloop()

if __name__ == "__main__":
    open_excel_data_entry()
